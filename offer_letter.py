from __future__ import annotations

import base64
from copy import deepcopy
from dataclasses import dataclass
from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from pathlib import Path
import re

import streamlit as st
from docx import Document


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "Offer Letter Template.docx"
LOGO_PATH = BASE_DIR / "Midea.png"
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
WORK_LOCATION_OPTIONS = {
    "Louisville": {
        "address": "2700 Chestnut Station Ct, Louisville, KY 40299",
        "location": "Louisville, KY",
    },
    "Boston": {
        "address": "260 Charles Street, Suite 401, Waltham, MA 02453",
        "location": "Boston, MA",
    },
}
JOB_SUMMARY_BY_TITLE: dict[str, str] = {
    "Mechanical Engineering": (
        "As a key role for introducing new US product platforms, the Mechanical Engineering "
        "Co-op will join a team of experts meeting the needs of the consumer American "
        "Consumer in one of our various departments. The Co-op will work on product "
        "development of new and innovative features for integration into a US product line. "
        "The Co-op will support the development of new technologies and designs which "
        "provide market-leading performance and consumer acceptance. The Co-op will also "
        "have the opportunity to assist in the development of Midea's US patent portfolio "
        "development, as well as actively working on individual part design, prototyping, "
        "and a wide variety of laboratory work in proving out conceptual designs."
    ),
    "Industrial Design": (
        "As a key role for the introduction of new US product platforms, the Industrial "
        "Design Co-op will join a team of experts meeting the needs of the consumer "
        "American Consumer in appliance technology. Midea's ID interns work hands-on with "
        "our full-time engineers to brainstorm, sketch, conceptualize, prototype, and/or "
        "add to the Midea patent portfolio. The intern will work on product development "
        "and unique feature creation with innovative graphics, colors, and product styling "
        "for integration into a US product line. The intern will support development of "
        "creative design elements that meet modern styling and branding of other Midea "
        "products. The intern will assist in the development of new technologies and "
        "designs which provide market-leading form and aesthetics. The intern may also have "
        "the opportunity to assist in the development of Midea's US patent portfolio "
        "development and assist with prototyping and laboratory work in proving out "
        "conceptual and aesthetic design."
    ),
    "Electrical Engineering": (
        "As a key role for the introduction of new US product platforms, the Electrical "
        "Engineering Co-op will support the introduction of new U.S. product platforms. "
        "The Co-op will work closely with the engineering team to design, develop, and test "
        "embedded systems for innovative home appliances. Key responsibilities include "
        "firmware configuration and optimization, implementing Over-the-Air (OTA) update "
        "functionality, and contributing to system architecture and design reviews. The "
        "Co-op will also collaborate with colleagues on circuit design, prototyping, PCB "
        "soldering, and debugging activities. Additional opportunities may involve "
        "supporting patent portfolio development and assisting in the creation of product "
        "specifications to guide R&D. This role offers hands-on experience with cutting-edge "
        "technologies while contributing to Midea's mission of delivering market-leading "
        "performance and consumer-focused innovations."
    ),
    "Computer Engineering": (
        "As a key role for the introduction of new US product platforms, the Computer "
        "Engineering Co-op will work alongside the engineering team to design, develop, and "
        "test embedded systems across a range of applications. This role may involve "
        "configuring and optimizing firmware for microprocessors, implementing secure "
        "Over-the-Air (OTA) updates, and assisting with hardware tasks such as PCB "
        "soldering and circuit debugging. The co-op will participate in design reviews, "
        "contribute to system architecture discussions, and maintain thorough documentation "
        "of specifications, testing procedures, and results to support collaboration and "
        "knowledge sharing."
    ),
    "Consumer and Marketing Insights (CMI)": (
        "As a key role for supporting consumer insights in new product development, the "
        "Consumer and Marketing Insights (\"CMI\") Co-op will join a team of specialists "
        "dedicated to understanding the needs of the target consumer. The Intern will "
        "support the Consumer Insights Specialist throughout all phases of consumer insights "
        "research. The Intern will help perform both quantitative and qualitative research "
        "to gain a clear picture of consumer behavior, preferences, and market trends. The "
        "Intern will assist in designing and developing questionnaires and surveys for "
        "various research projects, and help manage and maintain research equipment and "
        "materials. The Intern will ensure accuracy in data analysis and reporting, and "
        "assist in the development of detailed reports and presentations to share findings "
        "and recommendations with stakeholders. The Intern will also have the opportunity "
        "to utilize consumer insights to help guide the development of products to meet the "
        "needs of the target consumer."
    ),
    "Computer Science": (
        "As a key role for leveraging data and software to drive product development, the "
        "Computer Science Co-op will join a team of experts dedicated to delivering "
        "exceptional products for the US market. The Co-op will utilize Python to gather "
        "and build data models via web scraping, and will add functionalities, fix bugs in "
        "the code, gather data trends, and condense data into a presentable format for team "
        "members. The Co-op will work on product specifications containing adequate "
        "information to guide R&D activities. The Co-op may work on individual circuit "
        "designs, prototyping, and laboratory work in proving out conceptual designs. The "
        "Co-op will support the development of new technologies and designs which provide "
        "market-leading performance and consumer acceptance. The Co-op will also have the "
        "opportunity to work closely with the engineering team to design, develop, and test "
        "embedded systems for various applications."
    ),
    "Data Science/Data Analyst": (
        "As a key role for extracting and applying insights from data to drive product and "
        "business decisions, the Data Science/Data Analyst Co-op will join a team focused "
        "on leveraging data for competitive advantage. The Co-op will use various machine "
        "learning software to mine online product reviews, including competitive data, and "
        "apply machine learning techniques to determine sentiment and categorize feedback by "
        "product family. The Co-op will utilize findings to provide actionable feedback to "
        "R&D teams for product improvements and future planning. The role will also involve "
        "using machine learning to gather product insights from social platforms and "
        "building dashboards to democratize these insights for the team. The Co-op will "
        "assist in analyzing data from research events to recognize key insights and work "
        "with cross-functional teams to build sales tools from large datasets, covering "
        "areas such as sales incentives, competitive analysis, and service data. The Co-op "
        "will proactively look for opportunities to use their skills to improve team "
        "processes and will have the ability to meet with customers to discuss product "
        "reviews and ratings, helping to answer questions about their processes."
    ),
    "HR/IT": (
        "As a key role for providing integrated technical and human resources support, the "
        "HR/IT Intern will join a team dedicated to enhancing operational efficiency and "
        "employee experience. The Intern will provide hands-on technical support by "
        "troubleshooting hardware, software, and network issues, while assisting employees "
        "with system setup, configuration, and documentation. The Intern will respond to "
        "help desk tickets, document resolutions, and support data collection and analysis "
        "for IT improvement projects. The role will involve documenting IT and HR "
        "workflows, streamlining processes, and contributing ideas for improving efficiency "
        "and communication, including the creation of user guides and system logs. The "
        "Intern will collaborate with teams across departments to understand and support "
        "technology and HR needs. The Intern will participate in IT infrastructure or "
        "software upgrade projects and HR performance management initiatives under the "
        "guidance of department mentors. The Intern will also assist with full-cycle hiring "
        "activities, including screening resumes, scheduling interviews, and maintaining "
        "candidate records. Furthermore, the Intern will support employee engagement "
        "initiatives, help foster a positive company culture, and assist with the "
        "coordination of the co-op program, company events, and HR-led activities."
    ),
    "HR": (
        "As a key role for supporting the human resources function across the employee "
        "lifecycle, the HR Intern will join a team dedicated to fostering a positive and "
        "efficient workplace. The Intern will assist in the coordination of full-cycle "
        "hiring efforts, including screening resumes, scheduling candidates, and "
        "maintaining candidate records. The Intern will assist the HR manager in "
        "scheduling and facilitating a smooth new hire onboarding process. The role will "
        "involve supporting employee engagement initiatives and contributing to a positive "
        "company culture. The Intern will maintain accurate HR records and documentation to "
        "ensure data integrity and compliance. The Intern will also assist in the MARC "
        "co-op program coordination and event planning, and help with employee performance "
        "management activities under the guidance of the HR manager. The Intern will be "
        "assigned other duties as needed to support the HR department."
    ),
    "Materials Science and Engineering": (
        "As the key role for advancing materials innovation and supporting product "
        "development, the Materials Science and Engineering Co-op will join a team of "
        "experts dedicated to creating market-leading technologies. The Co-op will research "
        "cutting-edge engineering topics and assess potential innovation opportunities. The "
        "Co-op will plan and carry out laboratory experiments to support rapid prototyping "
        "projects, utilizing analytical equipment to test new materials concepts and "
        "processes. The role involves analyzing data and developing presentation materials "
        "for effective communication of results. The Co-op will prototype design concepts, "
        "assess performance, and summarize findings. The Co-op will interact with design "
        "engineers to discuss experimental results and conduct feasibility studies to aid "
        "with design innovation and product development. The Co-op will actively "
        "participate in project reviews and present results to cross-functional design "
        "teams. The Co-op will also have the opportunity to work with engineers on the "
        "development of new intellectual property."
    ),
    "IoT/Software Engineering": (
        "As the key role for advancing smart home technology and IoT solutions, the "
        "IoT/Software Engineering Co-op will join a team of experts dedicated to enhancing "
        "the Smart Home ecosystem. The Co-op will work with Senior Staff IoT and Software "
        "Engineers to implement new features into the Smart Home ecosystem. The Co-op will "
        "work on advanced development IoT projects pertaining to Residential and Central "
        "Air Conditioning products, such as creating a proof of concept to adjust an AC's "
        "setpoint temperature automatically based on a person's body temperature or room "
        "occupancy. The role involves working in cross-functional and cross-cultural teams "
        "to ensure IoT/smart home features and services conform to specifications and meet "
        "performance needs. The Co-op will assist in testing IoT devices to ensure they "
        "adhere to standards. The Co-op will also support other team members in their "
        "technical specialty as needed."
    ),
}
JOB_TITLE_ALIASES = {
    "mechanical engineering co-op": "Mechanical Engineering",
    "mechanical engineering coop": "Mechanical Engineering",
    "mechanical": "Mechanical Engineering",
    "materials science": "Materials Science and Engineering",
    "materials science co-op": "Materials Science and Engineering",
    "materials science and engineering co-op": "Materials Science and Engineering",
    "chemistry": "Materials Science and Engineering",
    "chemistry co-op": "Materials Science and Engineering",
    "industrial design co-op": "Industrial Design",
    "electrical engineering co-op": "Electrical Engineering",
    "electrical": "Electrical Engineering",
    "computer engineering co-op": "Computer Engineering",
    "consumer and marketing insights": "Consumer and Marketing Insights (CMI)",
    "consumer and marketing insights co-op": "Consumer and Marketing Insights (CMI)",
    "cmi": "Consumer and Marketing Insights (CMI)",
    "cmi co-op": "Consumer and Marketing Insights (CMI)",
    "computer science co-op": "Computer Science",
    "data science": "Data Science/Data Analyst",
    "data science co-op": "Data Science/Data Analyst",
    "data analyst": "Data Science/Data Analyst",
    "data analyst co-op": "Data Science/Data Analyst",
    "data science/data analyst co-op": "Data Science/Data Analyst",
    "hr/it intern": "HR/IT",
    "hr it": "HR/IT",
    "hr it intern": "HR/IT",
    "hr/it co-op": "HR/IT",
    "human resources": "HR",
    "human resources intern": "HR",
    "hr intern": "HR",
    "hr co-op": "HR",
    "iot": "IoT/Software Engineering",
    "iot co-op": "IoT/Software Engineering",
    "software engineering": "IoT/Software Engineering",
    "software engineering co-op": "IoT/Software Engineering",
    "iot/software engineering co-op": "IoT/Software Engineering",
    "iot software engineering": "IoT/Software Engineering",
}


@dataclass
class OfferLetterData:
    candidate_name: str
    letter_date: date
    position_title: str
    job_summary: str
    manager_name: str
    manager_title: str
    work_location: str
    location: str
    employment_start_date: date
    employment_end_date: date
    hourly_rate: Decimal
    relocation_assistance: Decimal
    sign_on_bonus: Decimal
    output_stem: str


def format_long_date(value: date) -> str:
    return f"{value.strftime('%B')} {value.day}, {value.year}"


def format_money(value: Decimal) -> str:
    normalized = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{normalized:,.2f}"


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "", value).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    cleaned = re.sub(r"_+", "_", cleaned)
    return cleaned.strip("._") or "offer_letter"


def compute_overtime_rate(hourly_rate: Decimal) -> Decimal:
    return hourly_rate * Decimal("1.5")


def get_image_data_uri(image_path: Path) -> str:
    encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def build_default_output_stem(candidate_name: str) -> str:
    return sanitize_filename(f"{candidate_name} Offer Letter")


def normalize_job_title(value: str) -> str:
    normalized = re.sub(r"\bco[\s-]?op\b", "co-op", value, flags=re.IGNORECASE)
    return " ".join(normalized.split()).casefold()


def get_job_summary_for_title(position_title: str) -> str:
    normalized_title = normalize_job_title(position_title)
    normalized_title = normalize_job_title(JOB_TITLE_ALIASES.get(normalized_title, normalized_title))
    for saved_title, saved_summary in JOB_SUMMARY_BY_TITLE.items():
        if normalize_job_title(saved_title) == normalized_title:
            return saved_summary
    return ""


def clear_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def add_styled_run(paragraph, text: str, source_run=None) -> None:
    run = paragraph.add_run(text)
    if source_run is not None and source_run._r.rPr is not None:
        run._r.insert(0, deepcopy(source_run._r.rPr))


def replace_paragraph_text(paragraph, text: str) -> None:
    original_runs = list(paragraph.runs)
    clear_runs(paragraph)

    if not text:
        paragraph.add_run("")
        return

    source_run = original_runs[0] if original_runs else None
    add_styled_run(paragraph, text, source_run)


def add_grouped_text_segments(paragraph, text: str, run_indexes: list[int], source_runs) -> None:
    if not text:
        return

    if not run_indexes:
        add_styled_run(paragraph, text, source_runs[0] if source_runs else None)
        return

    segment_start = 0
    current_run_index = run_indexes[0]
    for index, run_index in enumerate(run_indexes[1:], start=1):
        if run_index != current_run_index:
            add_styled_run(paragraph, text[segment_start:index], source_runs[current_run_index])
            segment_start = index
            current_run_index = run_index
    add_styled_run(paragraph, text[segment_start:], source_runs[current_run_index])


def iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def paragraph_has_placeholder(paragraph, placeholder_name: str) -> bool:
    pattern = r"\{\{\s*" + re.escape(placeholder_name) + r"\s*\}\}"
    return re.search(pattern, paragraph.text) is not None


def replace_placeholders(text: str, replacements: dict[str, str]) -> str:
    def lookup(match: re.Match[str]) -> str:
        key = match.group(1).strip()
        return replacements.get(key, match.group(0))

    return PLACEHOLDER_PATTERN.sub(lookup, text)


def replace_placeholders_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    original_runs = list(paragraph.runs)
    original_text = "".join(run.text for run in original_runs)
    if not PLACEHOLDER_PATTERN.search(original_text):
        return

    char_run_indexes: list[int] = []
    for run_index, run in enumerate(original_runs):
        char_run_indexes.extend([run_index] * len(run.text))

    segments: list[tuple[str, int | None, list[int] | None]] = []
    cursor = 0
    for match in PLACEHOLDER_PATTERN.finditer(original_text):
        if match.start() > cursor:
            plain_text = original_text[cursor:match.start()]
            plain_run_indexes = char_run_indexes[cursor:match.start()]
            segments.append((plain_text, None, plain_run_indexes))

        placeholder_key = match.group(1).strip()
        replacement_text = replacements.get(placeholder_key, match.group(0))
        placeholder_run_index = char_run_indexes[match.start()] if char_run_indexes else 0
        segments.append((replacement_text, placeholder_run_index, None))
        cursor = match.end()

    if cursor < len(original_text):
        trailing_text = original_text[cursor:]
        trailing_run_indexes = char_run_indexes[cursor:]
        segments.append((trailing_text, None, trailing_run_indexes))

    clear_runs(paragraph)
    for text, run_index, run_indexes in segments:
        if not text:
            continue
        if run_indexes is not None:
            add_grouped_text_segments(paragraph, text, run_indexes, original_runs)
        else:
            source_run = original_runs[run_index] if original_runs else None
            add_styled_run(paragraph, text, source_run)


def build_offer_letter_document(template_path: Path, data: OfferLetterData) -> Document:
    document = Document(template_path)

    if data.relocation_assistance == Decimal("0"):
        for paragraph in list(iter_paragraphs(document)):
            if paragraph_has_placeholder(paragraph, "relocation_fee"):
                remove_paragraph(paragraph)
    if data.sign_on_bonus == Decimal("0"):
        for paragraph in list(iter_paragraphs(document)):
            if paragraph_has_placeholder(paragraph, "sign_on"):
                remove_paragraph(paragraph)

    replacements = {
        "today_date": format_long_date(data.letter_date),
        "full_name": data.candidate_name,
        "job_title": data.position_title,
        "job_summary": data.job_summary,
        "manager_name": f" {data.manager_name}",
        "manager_title": data.manager_title,
        "work_location": data.work_location,
        "location": f"{data.location} ",
        "start_date": format_long_date(data.employment_start_date),
        "end_date": format_long_date(data.employment_end_date),
        "hourly_rate": format_money(data.hourly_rate),
        "overtime_rate": format_money(compute_overtime_rate(data.hourly_rate)),
        "relocation_fee": format_money(data.relocation_assistance),
        "sign_on": format_money(data.sign_on_bonus),
    }

    for paragraph in iter_paragraphs(document):
        replace_placeholders_in_paragraph(paragraph, replacements)

    return document


def build_offer_letter_bytes(template_path: Path, data: OfferLetterData) -> bytes:
    document = build_offer_letter_document(template_path, data)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_data(
    candidate_name: str,
    letter_date: date,
    position_title: str,
    job_summary: str,
    manager_name: str,
    manager_title: str,
    work_location: str,
    employment_start_date: date,
    employment_end_date: date,
    hourly_rate: float | int,
    relocation_assistance: float | int,
    sign_on_bonus: float | int,
    output_stem: str,
) -> OfferLetterData:
    location_label = work_location.strip()
    location_values = WORK_LOCATION_OPTIONS.get(
        location_label,
        {"address": location_label, "location": location_label},
    )
    return OfferLetterData(
        candidate_name=candidate_name.strip(),
        letter_date=letter_date,
        position_title=position_title.strip(),
        job_summary=job_summary.strip(),
        manager_name=manager_name.strip(),
        manager_title=manager_title.strip(),
        work_location=location_values["address"],
        location=location_values["location"],
        employment_start_date=employment_start_date,
        employment_end_date=employment_end_date,
        hourly_rate=Decimal(str(hourly_rate)),
        relocation_assistance=Decimal(str(relocation_assistance)),
        sign_on_bonus=Decimal(str(sign_on_bonus)),
        output_stem=sanitize_filename(output_stem),
    )


def main() -> None:
    st.set_page_config(page_title="MARC CO-OP Offer Letter Generator", layout="wide")

    logo_markup = ""
    if LOGO_PATH.exists():
        logo_markup = f'<img class="midea-logo" src="{get_image_data_uri(LOGO_PATH)}" alt="Midea logo" />'

    st.markdown(
        f"""
        <style>
        [data-testid="stAppViewContainer"] {{
            background: linear-gradient(180deg, #dff1ff 0%, #c7e4ff 100%);
        }}
        [data-testid="stHeader"] {{
            background: transparent;
        }}
        .main .block-container {{
            max-width: 1180px;
            padding-top: 3.5rem;
            padding-bottom: 2.5rem;
        }}
        div[data-testid="stVerticalBlockBorderWrapper"] {{
            background: rgba(255, 255, 255, 0.98);
            border: 1px solid rgba(104, 157, 214, 0.22);
            border-radius: 18px;
            padding: 1.25rem;
            box-shadow: 0 16px 36px rgba(72, 120, 174, 0.12);
        }}
        .midea-hero {{
            position: relative;
            min-height: 88px;
            margin-bottom: 1.75rem;
        }}
        .midea-logo {{
            position: absolute;
            top: 0;
            left: 0;
            width: 180px;
            height: auto;
        }}
        h1.marc-title {{
            text-align: center;
            color: #0f4d8a;
            margin: 0.75rem 0 2rem;
            letter-spacing: 0.02em;
        }}
        div[data-testid="stDownloadButton"] button {{
            width: 100%;
            min-height: 4rem;
            font-size: 1.18rem;
            font-weight: 600;
            border-radius: 12px;
        }}
        </style>
        <div class="midea-hero">
            {logo_markup}
            <h1 class="marc-title">MARC CO-OP Offer Letter Generator</h1>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if not TEMPLATE_PATH.exists():
        st.error(f"Template not found: {TEMPLATE_PATH}")
        return

    today = date.today()
    saved_job_titles = list(JOB_SUMMARY_BY_TITLE)
    job_title_options = saved_job_titles + ["Other"]

    with st.container(border=True):
        first_left, first_right = st.columns(2)
        with first_left:
            candidate_name = st.text_input("Full name", value="")
        with first_right:
            selected_job_title = st.selectbox(
                "Job title",
                options=job_title_options,
                index=None,
                placeholder="Select a job title",
            )

        position_title = selected_job_title or ""
        job_summary = ""

        if selected_job_title == "Other":
            position_title = st.text_input(
                "Custom job title (title only, no Co-op)",
                value="",
            )
            summary_title = st.selectbox(
                "Job summary template",
                options=saved_job_titles + ["Other"],
                index=None,
                placeholder="Select a job title for the summary",
            )
            if summary_title == "Other":
                job_summary = st.text_area(
                    "Job summary",
                    value="",
                    height=140,
                    placeholder="Paste the custom job summary here.",
                )
            else:
                job_summary = JOB_SUMMARY_BY_TITLE.get(summary_title or "", "")
        elif selected_job_title:
            job_summary = JOB_SUMMARY_BY_TITLE[selected_job_title]

        second_left, second_right = st.columns(2)
        with second_left:
            manager_name = st.text_input("Manager name", value="")
        with second_right:
            manager_title = st.text_input("Manager title", value="")

        third_left, third_right = st.columns(2)
        with third_left:
            employment_start_date = st.date_input("Employment start date", value=None)
        with third_right:
            employment_end_date = st.date_input("Employment end date", value=None)

        fourth_left, fourth_right = st.columns(2)
        with fourth_left:
            hourly_rate = st.number_input(
                "Hourly rate ($)",
                min_value=0.0,
                value=None,
                step=0.5,
                placeholder="Enter hourly rate",
            )
        with fourth_right:
            relocation_assistance = st.number_input(
                "Relocation assistance ($)",
                min_value=0.0,
                value=None,
                step=100.0,
                placeholder="Enter relocation assistance",
            )

        fifth_left, fifth_right = st.columns(2)
        with fifth_left:
            sign_on_bonus = st.number_input(
                "Sign-on bonus ($)",
                min_value=0.0,
                value=None,
                step=100.0,
                placeholder="Enter sign-on bonus",
            )
        with fifth_right:
            work_location = st.selectbox(
                "Working location",
                options=list(WORK_LOCATION_OPTIONS.keys()),
                index=None,
                placeholder="Select a location",
            )

        if relocation_assistance is None:
            relocation_assistance = 0.0
        if sign_on_bonus is None:
            sign_on_bonus = 0.0

        validation_errors = []
        if not candidate_name.strip():
            validation_errors.append("Name")
        if not (position_title or "").strip():
            validation_errors.append("Job title")
        if not manager_name.strip():
            validation_errors.append("Manager name")
        if not manager_title.strip():
            validation_errors.append("Manager title")
        if not job_summary.strip():
            validation_errors.append("Job summary")
        if not (work_location or "").strip():
            validation_errors.append("Working location")
        if employment_start_date is None or employment_end_date is None:
            validation_errors.append("Employment dates")
        if hourly_rate is None:
            validation_errors.append("Hourly rate")
        if (
            employment_start_date is not None
            and employment_end_date is not None
            and employment_end_date < employment_start_date
        ):
            validation_errors.append("Employment end date must be on or after the employment start date")

        final_output_stem = build_default_output_stem(candidate_name or "candidate")
        download_data = b""
        generation_error = None

        if not validation_errors:
            try:
                offer_data = build_data(
                    candidate_name=candidate_name,
                    letter_date=today,
                    position_title=position_title,
                    job_summary=job_summary,
                    manager_name=manager_name,
                    manager_title=manager_title,
                    work_location=work_location,
                    employment_start_date=employment_start_date,
                    employment_end_date=employment_end_date,
                    hourly_rate=hourly_rate,
                    relocation_assistance=relocation_assistance,
                    sign_on_bonus=sign_on_bonus,
                    output_stem=final_output_stem,
                )
                download_data = build_offer_letter_bytes(TEMPLATE_PATH, offer_data)
            except Exception as exc:
                generation_error = str(exc)

        if generation_error:
            st.error(f"Generation failed: {generation_error}")
        elif validation_errors:
            st.caption("Fill in all required fields before generating the letter.")

        button_left, button_center, button_right = st.columns([1.2, 1.6, 1.2])
        with button_center:
            st.download_button(
                label="Generate offer letter",
                data=download_data,
                file_name=f"{final_output_stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=bool(validation_errors or generation_error),
                use_container_width=True,
            )


if __name__ == "__main__":
    main()
